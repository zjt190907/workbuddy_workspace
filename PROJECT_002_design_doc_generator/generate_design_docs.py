# -*- coding: utf-8 -*-
"""
批量生成开发详细设计文档（增量模式）
基于模板 + 禅道需求数据 + SQL脚本
策略：使用 python-docx 打开模板，定位段落并修改内容，不删除范围

用法:
  python generate_design_docs.py              # 增量生成：只为尚未生成文档的需求生成
  python generate_design_docs.py --force       # 强制重新生成所有需求文档
  python generate_design_docs.py 2794 2866     # 只生成指定需求号的文档
  python generate_design_docs.py --force 2794  # 强制重新生成指定需求号
  python generate_design_docs.py --list        # 列出所有需求目录及文档生成状态
"""

import os, sys, json, re, traceback, argparse
from copy import deepcopy
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ============== 配置 ==============
BASE_DIR = r"D:\学习\2026"  # 需求目录所在根目录（SQL和输出文档的位置）
TEMPLATE_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "开发详细设计-模板_converted.docx")
ZENTAO_CACHE = os.path.join(os.path.dirname(os.path.abspath(__file__)), "zentao_stories_full.json")
ZENTAO_URL = "http://172.16.1.138:8080/zentao"
ZENTAO_ACCOUNT = "zhr"
ZENTAO_PASSWORD = "7890-poi"

# 目录名中提取需求号的正则（匹配全角/半角括号前、或目录名开头的数字）
_DIR_ID_PATTERN = re.compile(r'^(?:需求)?(\d{3,5})')


def scan_requirement_dirs():
    """自动扫描 BASE_DIR 下的需求目录，返回 [(目录名, 需求号), ...]
    规则：目录名以数字开头（可带"需求"前缀），且是子目录
    排除: .workbuddy、不以数字开头的目录
    """
    results = []
    if not os.path.isdir(BASE_DIR):
        return results
    for name in sorted(os.listdir(BASE_DIR)):
        full_path = os.path.join(BASE_DIR, name)
        if not os.path.isdir(full_path):
            continue
        if name.startswith(".") or name.startswith("_"):
            continue
        m = _DIR_ID_PATTERN.match(name)
        if m:
            results.append((name, m.group(1)))
    return results


def get_output_path(req_dir, story_id):
    """获取输出文档路径"""
    return os.path.join(BASE_DIR, req_dir, f"开发详细设计-需求{story_id}.docx")


def doc_exists(req_dir, story_id):
    """检查文档是否已存在"""
    return os.path.isfile(get_output_path(req_dir, story_id))


# ============== 禅道数据 ==============
def load_zentao_data(story_ids=None):
    """加载禅道需求数据。优先读缓存，缺失的从API补充获取。
    story_ids: 需要加载的需求号列表，None则从缓存全量加载
    """
    # 读取已有缓存
    cached = {}
    if os.path.exists(ZENTAO_CACHE):
        with open(ZENTAO_CACHE, "r", encoding="utf-8") as f:
            cached = json.load(f)

    # 确定需要从API获取的需求号
    if story_ids is None:
        # 没指定则用缓存全部
        if cached:
            return cached
        # 无缓存则扫描目录获取
        story_ids = [sid for _, sid in scan_requirement_dirs()]

    missing_ids = [sid for sid in story_ids if sid not in cached]

    if not missing_ids:
        # 全部命中缓存，只返回需要的
        if story_ids:
            return {sid: cached[sid] for sid in story_ids if sid in cached}
        return cached

    # 从禅道API获取缺失的
    import requests
    session = requests.Session()
    session.post(f"{ZENTAO_URL}/user-login.html",
                 data={"account": ZENTAO_ACCOUNT, "password": ZENTAO_PASSWORD}, timeout=10)

    for story_id in missing_ids:
        try:
            resp = session.get(f"{ZENTAO_URL}/story-view-{story_id}.json", timeout=10)
            if resp.status_code == 200:
                data = resp.json()
                if data.get("status") == "success":
                    sd = json.loads(data["data"]) if isinstance(data["data"], str) else data["data"]
                    story = sd.get("story", {})
                    cached[story_id] = {"id": story.get("id"), "title": story.get("title", ""),
                                        "spec": story.get("spec", ""), "verify": story.get("verify", "")}
                    print(f"  从禅道获取 Story {story_id}: {story.get('title', '')[:40]}")
                else:
                    print(f"  Story {story_id}: 禅道返回 status={data.get('status')}，可能需求号不存在")
            else:
                print(f"  Story {story_id}: HTTP {resp.status_code}")
        except Exception as e:
            print(f"  获取Story {story_id}失败: {e}")

    # 更新缓存文件
    with open(ZENTAO_CACHE, "w", encoding="utf-8") as f:
        json.dump(cached, f, ensure_ascii=False, indent=2)
    print(f"  禅道缓存已更新 ({len(cached)} 条)")

    if story_ids:
        return {sid: cached[sid] for sid in story_ids if sid in cached}
    return cached


def parse_spec_html(html_text):
    if not html_text:
        return {"background": "", "goal": "", "menus": [], "gap": ""}
    from bs4 import BeautifulSoup
    soup = BeautifulSoup(html_text, "html.parser")
    result = {"background": "", "goal": "", "menus": [], "gap": ""}
    all_tds = soup.find_all("td")
    for i, td in enumerate(all_tds):
        text = td.get_text(strip=True)
        if text == "需求背景":
            c = _get_content(td, all_tds, i)
            if c: result["background"] = c
        elif text == "需求目标":
            c = _get_content(td, all_tds, i)
            if c: result["goal"] = c
        elif text == "影响菜单":
            c = _get_content(td, all_tds, i)
            if c: result["menus"] = [m.strip() for m in re.split(r'[、，,\n]', c) if m.strip() and len(m.strip()) > 1]
        elif text == "功能差距":
            c = _get_content(td, all_tds, i)
            if c and c not in ("-", "—", "无"): result["gap"] = c
    return result

def _get_content(label_td, all_tds, idx):
    parent_tr = label_td.find_parent("tr")
    if parent_tr:
        tds = parent_tr.find_all("td")
        for j, td in enumerate(tds):
            if td == label_td:
                for k in range(j+1, len(tds)):
                    t = tds[k].get_text(strip=True)
                    if t and t not in ("需求背景","需求目标","影响菜单","功能差距","需求描述","功能定位","不受理影响","需求类型"):
                        return t
    if idx + 1 < len(all_tds):
        t = all_tds[idx+1].get_text(strip=True)
        if t and t not in ("需求背景","需求目标","影响菜单","功能差距","需求描述","功能定位","不受理影响","需求类型"):
            return t
    return ""


# ============== SQL 解析 ==============
def scan_sql_files(directory):
    structure_sqls, data_sqls = [], []
    if not os.path.isdir(directory): return structure_sqls, data_sqls
    for f in os.listdir(directory):
        if not f.endswith(".sql"): continue
        fp = os.path.join(directory, f)
        if "表数据" in f: data_sqls.append(fp)
        else: structure_sqls.append(fp)
    return structure_sqls, data_sqls

def read_sql_file(filepath):
    for enc in ["utf-8", "gbk", "latin1"]:
        try:
            with open(filepath, "r", encoding=enc) as f: return f.read()
        except: continue
    return ""

def parse_structure_sql(sql_text):
    result = {"create_tables": [], "alter_add": [], "alter_modify": [], "alter_drop": [], "create_index": []}
    # CREATE TABLE
    for m in re.finditer(r'CREATE\s+TABLE\s+(?:IF\s+NOT\s+EXISTS\s+)?`?(\w+)`?\.`?(\w+)`?\s*\((.*?)\)\s*ENGINE', sql_text, re.I|re.S):
        schema, tbl, body = m.group(1), m.group(2), m.group(3)
        cm = re.search(r"COMMENT\s*=\s*'([^']*)'", sql_text[m.start():m.end()+200], re.I)
        cols = []
        for line in body.split("\n"):
            line = line.strip().rstrip(",")
            if not line or line.upper().startswith(("PRIMARY","KEY","UNIQUE")): continue
            c = re.match(r'`(\w+)`\s+([\w]+(?:\([\d,\s]+\))?)'
                        r'(?:\s+CHARACTER\s+SET\s+\w+\s+COLLATE\s+\w+)?'
                        r'(?:\s+(NOT\s+NULL|NULL))?'
                        r'(?:\s+DEFAULT\s+(?:\S+|\'[^\']*\'))?'
                        r'\s*(?:COMMENT\s+\'([^\']*)\')?', line, re.I)
            if c:
                cols.append({"name":c.group(1),"type":c.group(2),
                    "nullable":"NOT NULL" if c.group(3) and "NOT NULL" in c.group(3).upper() else "NULL",
                    "default":"","comment":c.group(4) or ""})
        result["create_tables"].append({"schema":schema,"table":tbl,"comment":cm.group(1) if cm else "","columns":cols})

    # ALTER ADD
    for m in re.finditer(r'ALTER\s+TABLE\s+`?(\w+)`?\.`?(\w+)`?\s+ADD\s+(?:COLUMN\s+)?`?(\w+)`?\s+([\w()]+(?:\([\d,\s]+\))?)'
                        r'(?:\s+CHARACTER\s+SET\s+\w+\s+COLLATE\s+\w+)?'
                        r'(?:\s+(?:DEFAULT\s+\S+))?\s*(?:NULL|NOT\s+NULL)?'
                        r'\s*(?:COMMENT\s+[\'"]([^\'"]*)[\'"])?', sql_text, re.I):
        result["alter_add"].append({"schema":m.group(1),"table":m.group(2),"column":m.group(3),"type":m.group(4),"comment":m.group(5) or ""})

    # ALTER MODIFY
    for m in re.finditer(r'ALTER\s+TABLE\s+`?(\w+)`?\.`?(\w+)`?\s+MODIFY\s+COLUMN\s+`?(\w+)`?\s+([\w()]+(?:\([\d,\s]+\))?)'
                        r'(?:\s+CHARACTER\s+SET\s+\w+\s+COLLATE\s+\w+)?'
                        r'(?:\s+(?:DEFAULT\s+\S+))?\s*(?:NULL|NOT\s+NULL)?'
                        r'\s*(?:COMMENT\s+[\'"]([^\'"]*)[\'"])?', sql_text, re.I):
        result["alter_modify"].append({"schema":m.group(1),"table":m.group(2),"column":m.group(3),"type":m.group(4),"comment":m.group(5) or ""})

    # ALTER DROP
    for m in re.finditer(r'ALTER\s+TABLE\s+`?(\w+)`?\.`?(\w+)`?\s+DROP\s+COLUMN\s+`?(\w+)`?', sql_text, re.I):
        result["alter_drop"].append({"schema":m.group(1),"table":m.group(2),"column":m.group(3)})

    # CREATE INDEX
    for m in re.finditer(r'CREATE\s+INDEX\s+`?(\w+)`?\s+USING\s+BTREE\s+ON\s+`?(\w+)`?\.`?(\w+)`?\s*\(([^)]+)\)', sql_text, re.I):
        result["create_index"].append({"index_name":m.group(1),"schema":m.group(2),"table":m.group(3),"columns":m.group(4)})

    return result

def parse_data_sql(sql_text):
    result = {"updates": [], "inserts": [], "deletes": []}
    for stmt in re.split(r';\s*\n', sql_text):
        stmt = stmt.strip()
        if not stmt or stmt.startswith("--"): continue
        if re.match(r'UPDATE\s+', stmt, re.I):
            m = re.match(r'UPDATE\s+`?(\w+)`?\.`?(\w+)`?', stmt, re.I)
            result["updates"].append({"table": f"{m.group(1)}.{m.group(2)}" if m else "?", "sql": stmt[:500]})
        elif re.match(r'INSERT\s+', stmt, re.I):
            m = re.match(r'INSERT\s+INTO\s+`?(\w+)`?\.`?(\w+)`?', stmt, re.I)
            result["inserts"].append({"table": f"{m.group(1)}.{m.group(2)}" if m else "?", "sql": stmt[:500]})
        elif re.match(r'DELETE\s+', stmt, re.I):
            m = re.match(r'DELETE\s+FROM\s+`?(\w+)`?\.`?(\w+)`?', stmt, re.I)
            result["deletes"].append({"table": f"{m.group(1)}.{m.group(2)}" if m else "?", "sql": stmt[:500]})
    return result


# ============== Word 生成 ==============

def _add_run_to_para(para, text, font_name="仿宋", font_size=Pt(12), bold=False):
    """添加一个run到段落"""
    run = para.add_run(text)
    run.font.name = font_name
    run.font.size = font_size
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    return run

def _insert_para_after_elem(ref_elem, text, style_name="Body Text", font_name="仿宋", font_size=Pt(12), bold=False):
    """在指定元素后插入新段落，返回新段落元素"""
    new_p = parse_xml(f'<w:p {nsdecls("w")}><w:pPr><w:pStyle w:val="{style_name}"/></w:pPr></w:p>')
    ref_elem.addnext(new_p)
    # 创建Paragraph包装
    from docx.text.paragraph import Paragraph
    para = Paragraph(new_p, new_p.getparent())
    _add_run_to_para(para, text, font_name, font_size, bold)
    return new_p

def _insert_table_after_elem(ref_elem, headers, rows):
    """在指定元素后插入表格，返回表格元素"""
    num_rows = len(rows) + 1
    num_cols = len(headers)
    # 创建表格XML
    tbl_xml = f'<w:tbl {nsdecls("w")}>'
    tbl_xml += '<w:tblPr><w:tblStyle w:val="TableGrid"/><w:tblW w:w="0" w:type="auto"/>'
    tbl_xml += '<w:tblBorders>'
    for border in ['top','left','bottom','right','insideH','insideV']:
        tbl_xml += f'<w:{border} w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
    tbl_xml += '</w:tblBorders></w:tblPr>'
    tbl_xml += '<w:tblGrid>'
    for _ in range(num_cols):
        tbl_xml += '<w:gridCol w:w="1800"/>'
    tbl_xml += '</w:tblGrid>'

    # Header row
    tbl_xml += '<w:tr>'
    for h in headers:
        tbl_xml += f'<w:tc><w:tcPr><w:shd w:val="clear" w:color="auto" w:fill="D9E2F3"/></w:tcPr><w:p><w:pPr><w:rPr><w:b/><w:rFonts w:eastAsia="黑体"/><w:sz w:val="20"/></w:rPr></w:pPr><w:r><w:rPr><w:b/><w:rFonts w:eastAsia="黑体"/><w:sz w:val="20"/></w:rPr><w:t>{h}</w:t></w:r></w:p></w:tc>'
    tbl_xml += '</w:tr>'

    # Data rows
    for row in rows:
        tbl_xml += '<w:tr>'
        for cell in row:
            cell_escaped = str(cell).replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            tbl_xml += f'<w:tc><w:p><w:pPr><w:rPr><w:rFonts w:eastAsia="仿宋"/><w:sz w:val="18"/></w:rPr></w:pPr><w:r><w:rPr><w:rFonts w:eastAsia="仿宋"/><w:sz w:val="18"/></w:rPr><w:t>{cell_escaped}</w:t></w:r></w:p></w:tc>'
        tbl_xml += '</w:tr>'

    tbl_xml += '</w:tbl>'
    tbl_elem = parse_xml(tbl_xml)
    ref_elem.addnext(tbl_elem)
    return tbl_elem


def _find_para_index(doc, style_name, text_match):
    """查找段落索引"""
    for i, para in enumerate(doc.paragraphs):
        if para.style.name == style_name and text_match in para.text:
            return i
    return None


def generate_doc(requirement_dir, story_id, zentao_data):
    """为单个需求生成开发详细设计文档"""
    req_path = os.path.join(BASE_DIR, requirement_dir)
    output_path = os.path.join(req_path, f"开发详细设计-需求{story_id}.docx")

    print(f"\n处理需求 {story_id}: {requirement_dir}")

    # 加载模板
    doc = Document(TEMPLATE_PATH)

    # 1. 修改标题
    for para in doc.paragraphs:
        if para.style.name == "Heading 2" and "开发详细设计" in para.text and "功能" not in para.text:
            for run in para.runs:
                if "开发详细设计" in run.text:
                    run.text = f"开发详细设计-需求{story_id}"
            break

    # 2. 获取数据
    story = zentao_data.get(story_id, {})
    spec_info = parse_spec_html(story.get("spec", ""))

    # 3. SQL
    structure_sqls, data_sqls = scan_sql_files(req_path)
    all_structure = {"create_tables":[], "alter_add":[], "alter_modify":[], "alter_drop":[], "create_index":[]}
    all_data = {"updates":[], "inserts":[], "deletes":[]}
    for sf in structure_sqls:
        p = parse_structure_sql(read_sql_file(sf))
        for k in all_structure: all_structure[k].extend(p[k])
    for sf in data_sqls:
        p = parse_data_sql(read_sql_file(sf))
        for k in all_data: all_data[k].extend(p[k])

    print(f"  SQL: 结构类={len(structure_sqls)}, 数据类={len(data_sqls)}")

    # 4. 填充 1.1 功能描述
    _fill_11(doc, spec_info, story)

    # 5. 填充 1.2 功能清单
    _fill_12(doc, spec_info)

    # 6. 填充 1.3.3 数据库设计
    _fill_133(doc, all_structure)

    # 7. 填充 1.7 旧数据处理
    _fill_17(doc, all_data)

    # 8. 保存
    doc.save(output_path)
    print(f"  [OK] 已生成: {output_path}")
    return True


def _fill_11(doc, spec_info, story):
    """填充1.1功能描述：修改功能描述标题后的空段落"""
    idx = _find_para_index(doc, "Heading 2", "功能描述")
    if idx is None: return

    # 构建描述文本
    background = spec_info.get("background", "")
    goal = spec_info.get("goal", "")
    gap = spec_info.get("gap", "")

    desc_text = ""
    if background: desc_text += f"需求背景：{background}"
    if goal:
        if desc_text: desc_text += "\n"
        desc_text += f"需求目标：{goal}"
    if gap and gap not in ("-", "—", "无"):
        if desc_text: desc_text += "\n"
        desc_text += f"功能差距：{gap}"
    if not desc_text:
        desc_text = story.get("title", "详见禅道需求")

    # 找到功能描述后的空段落并填充（每段一个run，换行用add_break）
    for i in range(idx + 1, min(idx + 5, len(doc.paragraphs))):
        para = doc.paragraphs[i]
        if para.style.name in ("Body Text", "Normal", "FinBodyStyle") and not para.text.strip():
            # 清空已有runs
            for run in para.runs:
                run.text = ""
            # 按换行分割，分段添加
            parts = desc_text.split("\n")
            for pi, part in enumerate(parts):
                if pi > 0:
                    para.add_run().add_break()
                _add_run_to_para(para, part, "仿宋", Pt(12))
            return

    # 如果没有空段落，在标题后插入
    heading_para = doc.paragraphs[idx]
    _insert_para_after_elem(heading_para._element, desc_text, "Body Text", "仿宋", Pt(12))


def _fill_12(doc, spec_info):
    """填充1.2功能清单：修改功能清单表格"""
    menus = spec_info.get("menus", [])
    if not menus: return

    for table in doc.tables:
        if len(table.rows) > 0 and "功能点归属" in table.rows[0].cells[0].text:
            # 填充第二行
            if len(table.rows) > 1:
                row = table.rows[1]
                row.cells[0].text = "本需求"
                row.cells[2].text = menus[0]
                row.cells[3].text = menus[0]
                row.cells[4].text = ""
            # 添加更多行
            for menu in menus[1:]:
                row = table.add_row()
                row.cells[0].text = "本需求"
                row.cells[2].text = menu
                row.cells[3].text = menu
                row.cells[4].text = ""
            break


def _fill_133(doc, structure_data):
    """填充1.3.3数据库设计"""
    has_content = any(structure_data[k] for k in structure_data)

    # 找到"数据库设计"标题
    db_idx = _find_para_index(doc, "Heading 3", "数据库设计")
    if db_idx is None: return

    db_para = doc.paragraphs[db_idx]

    if not has_content:
        # 找到数据库设计后最近的Body Text段落，设为"无"
        # 模板中数据库设计后没有Body Text，需要插入
        _insert_para_after_elem(db_para._element, "无", "Body Text", "仿宋", Pt(12))
        return

    # 收集要插入的元素（需要倒序插入，因为addnext总是在后面加）
    # 所以我们先按正序收集，然后倒序插入
    insert_list = []

    # 1. 新增表
    if structure_data["create_tables"]:
        insert_list.append(("heading", "一、新增表"))
        for ct in structure_data["create_tables"]:
            full_name = f"{ct['schema']}.{ct['table']}"
            comment = f"（{ct['comment']}）" if ct['comment'] else ""
            insert_list.append(("subheading", f"  {full_name}{comment}"))
            if ct["columns"]:
                # 表格数据
                headers = ["字段名", "类型", "为空", "默认值", "说明"]
                rows = [[c["name"], c["type"], c["nullable"], c.get("default",""), c.get("comment","")] for c in ct["columns"]]
                insert_list.append(("table", headers, rows))

    # 2. 字段变更
    has_alter = structure_data["alter_add"] or structure_data["alter_modify"] or structure_data["alter_drop"]
    if has_alter:
        sec_num = "二" if structure_data["create_tables"] else "一"
        insert_list.append(("heading", f"{sec_num}、字段变更"))

        table_groups = {}
        for item in structure_data["alter_add"]:
            key = f"{item['schema']}.{item['table']}"
            table_groups.setdefault(key, {"add":[],"modify":[],"drop":[]})["add"].append(item)
        for item in structure_data["alter_modify"]:
            key = f"{item['schema']}.{item['table']}"
            table_groups.setdefault(key, {"add":[],"modify":[],"drop":[]})["modify"].append(item)
        for item in structure_data["alter_drop"]:
            key = f"{item['schema']}.{item['table']}"
            table_groups.setdefault(key, {"add":[],"modify":[],"drop":[]})["drop"].append(item)

        for tname, changes in table_groups.items():
            if changes["add"]:
                insert_list.append(("text", f"  {tname} 新增字段："))
                for item in changes["add"]:
                    c = f" — {item['comment']}" if item['comment'] else ""
                    insert_list.append(("text", f"    · {item['column']} {item['type']}{c}"))
            if changes["modify"]:
                insert_list.append(("text", f"  {tname} 修改字段："))
                for item in changes["modify"]:
                    c = f" — {item['comment']}" if item['comment'] else ""
                    insert_list.append(("text", f"    · {item['column']} {item['type']}{c}"))
            if changes["drop"]:
                insert_list.append(("text", f"  {tname} 删除字段："))
                for item in changes["drop"]:
                    insert_list.append(("text", f"    · {item['column']}"))

    # 3. 索引变更
    if structure_data["create_index"]:
        sec_num = "一二三四五"[[structure_data["create_tables"], has_alter].count(True)]
        insert_list.append(("heading", f"{sec_num}、索引变更"))
        for idx in structure_data["create_index"]:
            insert_list.append(("text", f"  · {idx['index_name']} ON {idx['schema']}.{idx['table']} ({idx['columns']})"))

    # 正序插入（addnext + 更新ref_elem = 正序）
    ref_elem = db_para._element
    for item in insert_list:
        if item[0] == "heading":
            ref_elem = _insert_para_after_elem(ref_elem, item[1], "Body Text", "黑体", Pt(12), bold=True)
        elif item[0] == "subheading":
            ref_elem = _insert_para_after_elem(ref_elem, item[1], "Body Text", "仿宋", Pt(12), bold=True)
        elif item[0] == "text":
            ref_elem = _insert_para_after_elem(ref_elem, item[1], "Body Text", "仿宋", Pt(12))
        elif item[0] == "table":
            ref_elem = _insert_table_after_elem(ref_elem, item[1], item[2])


def _fill_17(doc, data_info):
    """填充1.7旧数据处理"""
    has_content = any(data_info[k] for k in data_info)

    old_idx = _find_para_index(doc, "Heading 2", "旧数据处理")
    if old_idx is None: return

    old_para = doc.paragraphs[old_idx]

    if not has_content:
        # 找到旧数据处理后的"无"字段落，确保存在
        # 模板中已有"无"，检查一下
        for i in range(old_idx + 1, min(old_idx + 5, len(doc.paragraphs))):
            if doc.paragraphs[i].style.name in ("Body Text",) and doc.paragraphs[i].text.strip() == "无":
                return  # 已有"无"
        _insert_para_after_elem(old_para._element, "无", "Body Text", "仿宋", Pt(12))
        return

    insert_list = []
    count = 0

    if data_info["updates"]:
        count += 1
        ch = "一二三四五六七"[count-1]
        insert_list.append(("heading", f"{ch}、数据修正"))
        for upd in data_info["updates"]:
            insert_list.append(("text", f"  · {upd['table']}"))
            sql_brief = upd['sql'][:200].replace('\n',' ').replace('\r','')
            insert_list.append(("sql", f"    {sql_brief}"))

    if data_info["inserts"]:
        count += 1
        ch = "一二三四五六七"[count-1]
        insert_list.append(("heading", f"{ch}、初始化数据"))
        for ins in data_info["inserts"]:
            insert_list.append(("text", f"  · {ins['table']}"))
            sql_brief = ins['sql'][:200].replace('\n',' ').replace('\r','')
            insert_list.append(("sql", f"    {sql_brief}"))

    if data_info["deletes"]:
        count += 1
        ch = "一二三四五六七"[count-1]
        insert_list.append(("heading", f"{ch}、数据清理"))
        for dele in data_info["deletes"]:
            insert_list.append(("text", f"  · {dele['table']}"))
            sql_brief = dele['sql'][:200].replace('\n',' ').replace('\r','')
            insert_list.append(("sql", f"    {sql_brief}"))

    ref_elem = old_para._element
    for item in insert_list:
        if item[0] == "heading":
            ref_elem = _insert_para_after_elem(ref_elem, item[1], "Body Text", "黑体", Pt(12), bold=True)
        elif item[0] == "text":
            ref_elem = _insert_para_after_elem(ref_elem, item[1], "Body Text", "仿宋", Pt(12))
        elif item[0] == "sql":
            ref_elem = _insert_para_after_elem(ref_elem, item[1], "Body Text", "Consolas", Pt(9))


# ============== 主程序 ==============
def list_status():
    """列出所有需求目录及文档生成状态"""
    all_dirs = scan_requirement_dirs()
    if not all_dirs:
        print("未发现需求目录")
        return

    print(f"{'需求号':<8} {'目录名':<45} {'状态'}")
    print("-" * 80)
    for req_dir, story_id in all_dirs:
        exists = doc_exists(req_dir, story_id)
        status = "已生成" if exists else "未生成"
        print(f"{story_id:<8} {req_dir:<45} {status}")

    generated = sum(1 for d, s in all_dirs if doc_exists(d, s))
    print(f"\n共 {len(all_dirs)} 个需求目录，已生成 {generated} 个，待生成 {len(all_dirs) - generated} 个")


def main():
    parser = argparse.ArgumentParser(
        description="开发详细设计文档批量生成（增量模式）",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例:
  python generate_design_docs.py              增量生成：只为尚未生成文档的需求生成
  python generate_design_docs.py --force       强制重新生成所有需求文档
  python generate_design_docs.py 2794 2866     只生成指定需求号的文档
  python generate_design_docs.py --force 2794  强制重新生成指定需求号
  python generate_design_docs.py --list        列出所有需求目录及文档生成状态
        """)
    parser.add_argument("--force", action="store_true", help="强制重新生成，即使文档已存在")
    parser.add_argument("--list", action="store_true", help="列出所有需求目录及生成状态，不生成文档")
    parser.add_argument("ids", nargs="*", help="指定需求号（如 2794 2866），只生成这些需求的文档")
    args = parser.parse_args()

    # 列表模式
    if args.list:
        list_status()
        return

    # 确定要处理的需求列表
    all_dirs = scan_requirement_dirs()
    if not all_dirs:
        print("未发现需求目录")
        return

    if args.ids:
        # 指定了需求号：过滤匹配的
        id_set = set(args.ids)
        targets = [(d, s) for d, s in all_dirs if s in id_set]
        not_found = id_set - {s for _, s in targets}
        if not_found:
            print(f"警告: 以下需求号未找到对应目录: {', '.join(sorted(not_found))}")
    else:
        targets = all_dirs

    # 增量过滤：跳过已存在的
    if not args.force:
        skipped = [(d, s) for d, s in targets if doc_exists(d, s)]
        targets = [(d, s) for d, s in targets if not doc_exists(d, s)]
        if skipped:
            print(f"跳过 {len(skipped)} 个已存在文档的需求（使用 --force 强制重新生成）")

    if not targets:
        print("没有需要生成的文档")
        return

    print("=" * 60)
    print(f"开发详细设计文档生成（{'强制' if args.force else '增量'}模式）")
    print(f"待处理: {len(targets)} 个需求")
    print("=" * 60)

    # 加载禅道数据
    print("\n[1/3] 加载禅道需求数据...")
    story_ids = [s for _, s in targets]
    zentao_data = load_zentao_data(story_ids)
    print(f"  已加载 {len(zentao_data)} 个需求数据")

    # 生成文档
    print("\n[2/3] 生成文档...")
    success = fail = skip = 0
    for req_dir, story_id in targets:
        try:
            generate_doc(req_dir, story_id, zentao_data)
            success += 1
        except Exception as e:
            print(f"  [FAIL] 需求{story_id}失败: {e}")
            traceback.print_exc()
            fail += 1

    print(f"\n[3/3] 完成！成功: {success}, 失败: {fail}")
    if fail > 0:
        print("  失败的需求可检查上方错误信息后重新运行")


if __name__ == "__main__":
    main()
